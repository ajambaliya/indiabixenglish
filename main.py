import logging
import asyncio
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
import urllib3
from telegram import Bot
from telegram.constants import PollType
from telegram.error import TelegramError
from datetime import datetime
import os
import pytz
import pymongo
from pymongo import MongoClient
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
import time
import subprocess

# Disable SSL/TLS-related warnings
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# Configuration from environment variables
MONGO_CONNECTION_STRING = os.environ.get('MONGO_CONNECTION_STRING')
TELEGRAM_BOT_TOKEN = os.environ.get('TELEGRAM_BOT_TOKEN')
TELEGRAM_CHAT_ID = os.environ.get('TELEGRAM_CHAT_ID')  # Use chat_id directly
TEMPLATE_URL = os.environ.get('TEMPLATE_URL')
TELEGRAM_CHANNEL_URL = "https://t.me/daily_current_all_source"  # Replace with your actual channel URL

# Validate that TELEGRAM_CHAT_ID is set and not empty
if not TELEGRAM_CHAT_ID:
    raise ValueError("TELEGRAM_CHAT_ID is not set or is empty.")

DB_NAME = 'IndiaBixEnglish'
COLLECTION_NAME = 'urls'

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TelegramQuizBot:
    def __init__(self, token, chat_id):
        self.bot = Bot(token=token)
        self.chat_id = chat_id

    def truncate_text(self, text, max_length):
        return text[:max_length-3] + '...' if len(text) > max_length else text

    async def send_poll(self, question_doc):
        # Append @daily_current_all_source to the question text
        question = self.truncate_text(question_doc["question"], 300) + " @daily_current_all_source"
        options = [self.truncate_text(opt, 100) for opt in question_doc["options"]]
        correct_option = question_doc["value_in_braces"]
        explanation = self.truncate_text(question_doc["explanation"], 200)

        option_mapping = {chr(65+i): i for i in range(len(options))}

        try:
            correct_option_id = option_mapping.get(correct_option)
            if correct_option_id is None:
                logger.error(f"Correct option '{correct_option}' not found in options: {options}")
                return

            # Send the poll with the updated question including @daily_current_all_source
            await self.bot.send_poll(
                chat_id=self.chat_id,
                question=question,
                options=options,
                is_anonymous=True,
                type=PollType.QUIZ,
                correct_option_id=correct_option_id,
                explanation=explanation
            )
            logger.info(f"Sent poll: {question}")

        except TelegramError as e:
            logger.error(f"Failed to send poll: {e.message}")
            logger.error(f"Full error details: {e}")

def get_current_month():
    ist = pytz.timezone('Asia/Kolkata')
    current_date = datetime.now(ist)
    return f"{current_date.month:02d}"

def connect_to_mongo():
    client = MongoClient(MONGO_CONNECTION_STRING)
    db = client[DB_NAME]
    collection = db[COLLECTION_NAME]
    return collection

# Gracefully handle documents without 'url' field
def get_scraped_urls(collection):
    urls = set()
    for doc in collection.find({}, {'url': 1}):
        url = doc.get('url')
        if url:
            urls.add(url)
        else:
            logger.warning(f"Document without 'url' field encountered: {doc}")
    return urls

# Clean up documents missing 'url' field
def clean_up_documents_without_url(collection):
    result = collection.delete_many({"url": {"$exists": False}})
    logger.info(f"Deleted {result.deleted_count} documents without 'url' field.")

def store_scraped_urls(collection, urls):
    for url in urls:
        collection.update_one({'url': url}, {'$set': {'url': url}}, upsert=True)

def scrape_latest_questions(latest_link):
    try:
        response = requests.get(latest_link, verify=False)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')

        question_docs = []

        question_divs = soup.find_all("div", class_="bix-div-container")

        for question_div in question_divs:
            try:
                qtxt = question_div.find("div", class_="bix-td-qtxt").text.strip()
                options_div = question_div.find("div", class_="bix-tbl-options")
                option_rows = options_div.find_all("div", class_="bix-opt-row")
                options = [option_row.find("div", class_="bix-td-option-val").text.strip() for option_row in option_rows]

                hidden_input = question_div.find("input", class_="jq-hdnakq")
                value_in_braces = hidden_input['value'].split('{', 1)[-1].rsplit('}', 1)[0] if hidden_input and 'value' in hidden_input.attrs else ""

                answer_div = question_div.find("div", class_="bix-div-answer")
                explanation = answer_div.find("div", class_="bix-ans-description").text.strip()

                question_doc = {
                    "question": qtxt,
                    "options": options,
                    "value_in_braces": value_in_braces,
                    "explanation": explanation
                }

                question_docs.append(question_doc)

            except Exception as e:
                logger.error(f"Error scraping content: {e}")

        return question_docs

    except requests.exceptions.RequestException as e:
        logger.error(f"Error fetching URL: {e}")
        return []

async def send_new_questions_to_telegram(new_questions):
    bot = TelegramQuizBot(TELEGRAM_BOT_TOKEN, TELEGRAM_CHAT_ID)
    for question in new_questions:
        await bot.send_poll(question)
        await asyncio.sleep(3)  # Rate limit to avoid spamming

def insert_content_from_top(doc, content_list):
    for content in content_list:
        new_para = doc.add_paragraph()
        run = new_para.add_run(content['text'])
        
        if content['type'] == 'question':
            run.bold = True
            new_para.paragraph_format.space_after = Pt(10)
            new_para.paragraph_format.space_before = Pt(10)
            run.font.size = Pt(14)
            
        elif content['type'] in ['options']:
            new_para.paragraph_format.left_indent = Pt(20)
            run.font.size = Pt(12)
            new_para.paragraph_format.space_after = Pt(6)
            
        elif content['type'] == 'answer':
            run.bold = True
            run.underline = True
            run.font.size = Pt(12)
            new_para.paragraph_format.space_before = Pt(10)
            
        elif content['type'] == 'explanation':
            run.italic = True
            run.font.size = Pt(12)
            new_para.paragraph_format.space_after = Pt(10)
            
        new_para.paragraph_format.line_spacing = 1.5

    add_promotional_message(doc)

def add_promotional_message(doc):
    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    run = para.add_run("📢 Join our Telegram Channel for daily quizzes and updates: ")
    run.bold = True
    run.font.size = Pt(14)
    
    url_run = para.add_run(f" {TELEGRAM_CHANNEL_URL}")
    url_run.font.color.rgb = RGBColor(0, 102, 204)
    url_run.font.size = Pt(19)
    url_run.bold = True

def prepare_content_list(question_docs):
    content_list = []
    for i, question in enumerate(question_docs, 1):
        content_list.extend([
            {'type': 'question', 'text': f"Question {i}: {question['question']}"},
            {'type': 'options', 'text': "Options:"},
            *[{'type': 'options', 'text': f"{chr(65+j)}. {opt}"} for j, opt in enumerate(question['options'])],
            {'type': 'answer', 'text': f"Correct Answer: {question['value_in_braces']}"},
            {'type': 'explanation', 'text': f"Explanation: {question['explanation']}"},
            {'type': 'space', 'text': "\n"}
        ])
    return content_list

def download_template(url):
    download_url = url.replace('/edit?usp=sharing', '/export?format=docx')
    try:
        response = requests.get(download_url)
        response.raise_for_status()
        return io.BytesIO(response.content)
    except requests.exceptions.RequestException as e:
        logger.error(f"Error downloading template: {e}")
        raise

def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        output_dir = os.path.dirname(pdf_path)
        result = subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path], 
                                check=True, capture_output=True, text=True)
        logger.info(f"LibreOffice conversion output: {result.stdout}")
        logger.error(f"LibreOffice conversion error output: {result.stderr}")
        
        pdf_temp_path = os.path.join(output_dir, os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
        if os.path.exists(pdf_temp_path):
            os.rename(pdf_temp_path, pdf_path)
            logger.info(f"Successfully converted DOCX to PDF: {pdf_path}")
        else:
            raise FileNotFoundError(f"PDF file not found at expected location: {pdf_temp_path}")
    except subprocess.CalledProcessError as e:
        logger.error(f"LibreOffice conversion failed: {e}")
        raise
    except Exception as e:
        logger.error(f"Error converting DOCX to PDF: {e}")
        raise

async def send_pdf_to_telegram(bot, chat_id, pdf_path, caption):
    try:
        with open(pdf_path, 'rb') as pdf_file:
            await bot.send_document(
                chat_id=chat_id,
                document=pdf_file,
                caption=caption
            )
        logger.info(f"Sent PDF to chat_id: {chat_id}")
    except TelegramError as e:
        logger.error(f"Failed to send PDF: {e.message}")

# Enhanced caption with symbols and design elements
def generate_pdf_caption(quiz_date, question_count):
    return (
        f"🎯 Current Affairs Quiz - {quiz_date}\n\n"
        f"📝 PDF Contents:\n"
        f"• Total Questions: {question_count}\n\n"
        f"🔍 Boost Your Knowledge: Stay updated with daily quizzes and enhance your knowledge!\n\n"
        f"💡 For More Quizzes:\n"
        f"Join our Telegram channel @Daily_Current_All_Source to receive daily updates.\n"
        f"🚀 Stay Ahead, Stay Informed!\n\n"
        f"⚡️ Test your knowledge today! 📊"
    )

def extract_date_from_url(url):
    parts = url.split("/")
    try:
        date_str = parts[-2]
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        return date_obj.strftime("%d %B %Y")
    except (ValueError, IndexError):
        logger.warning(f"Date extraction failed for URL: {url}")
        return datetime.now().strftime("%d %B %Y")

async def main():
    collection = connect_to_mongo()
    clean_up_documents_without_url(collection)  # Clean up documents without 'url' field (optional)
    stored_urls = get_scraped_urls(collection)
    
    url = "https://www.indiabix.com/current-affairs/questions-and-answers/"
    month_digit = get_current_month()

    response = requests.get(url, verify=False)
    response.raise_for_status()
    soup = BeautifulSoup(response.text, 'html.parser')
    link_elements = soup.find_all("a", class_="text-link me-3")

    valid_links = []
    for link_element in link_elements:
        href = link_element.get("href")
        if f"/current-affairs/2024-{month_digit}-" in href:
            full_url = urljoin("https://www.indiabix.com/", href)
            if full_url not in stored_urls:
                valid_links.append(full_url)

    if not valid_links:
        logger.info("No new valid links found.")
        return

    valid_links.sort(key=lambda x: datetime.strptime(x.split("/")[-2], "%Y-%m-%d"))

    for link in valid_links:
        logger.info(f"Scraping link: {link}")

        question_docs = scrape_latest_questions(link)
        
        if question_docs:
            store_scraped_urls(collection, [link])
            await send_new_questions_to_telegram(question_docs)

            content_list = prepare_content_list(question_docs)

            template_bytes = download_template(TEMPLATE_URL)
            doc = Document(template_bytes)
            insert_content_from_top(doc, content_list)

            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                doc.save(tmp_docx.name)

            pdf_filename = f"current_affairs_{datetime.now().strftime('%Y%m%d')}.pdf"
            pdf_path = os.path.abspath(pdf_filename)
            convert_docx_to_pdf(os.path.abspath(tmp_docx.name), pdf_path)

            quiz_date = extract_date_from_url(link)

            bot = Bot(token=TELEGRAM_BOT_TOKEN)
            caption = generate_pdf_caption(quiz_date, len(question_docs))
            await send_pdf_to_telegram(bot, TELEGRAM_CHAT_ID, pdf_path, caption)

            os.unlink(tmp_docx.name)
            os.remove(pdf_path)
        
        else:
            logger.info(f"No questions found for link: {link}")
        
        time.sleep(5)

if __name__ == "__main__":
    asyncio.run(main())
